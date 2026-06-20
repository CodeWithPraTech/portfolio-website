from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("Statistics_Day_5_Discrete_Distributions_Notes.docx")


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 89, 89)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")


def style_run(run, bold=False, italic=False, color=None, size=None):
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    p.add_run(text)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.style = "Normal"
    p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    p = cell.paragraphs[0]
    run = p.add_run(title)
    style_run(run, bold=True, color=DARK_BLUE)
    p.add_run(f" {body}")
    doc.add_paragraph()


def add_formula(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(20, 38, 60)
    return p


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, h in enumerate(headers):
        hdr[idx].text = h
        set_cell_shading(hdr[idx], LIGHT_BLUE)
        for paragraph in hdr[idx].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = DARK_BLUE
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    set_table_width(table, widths)
    doc.add_paragraph()
    return table


def configure_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def build_doc():
    doc = Document()
    configure_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Statistics Day 5 Notes: Discrete Distributions")
    style_run(run, bold=True, color=RGBColor(11, 37, 69), size=22)
    subtitle = doc.add_paragraph()
    subtitle.add_run("Bernoulli, Binomial, Geometric, and Poisson | MAANG + BFSI + Research Engineer depth")
    subtitle.runs[0].font.color.rgb = MUTED

    add_callout(
        doc,
        "Core idea:",
        "A distribution is a model for uncertainty. In interviews and real work, the hard part is not memorizing formulas; it is matching the business process to the right assumptions."
    )

    add_heading(doc, "1. What This Session Covers", 1)
    add_bullets(doc, [
        "Bernoulli distribution: one trial with success/failure outcome.",
        "Binomial distribution: count successes in a fixed number of independent Bernoulli trials.",
        "Geometric distribution: waiting time until the first success.",
        "Poisson distribution: count rare events in a fixed interval of time, space, traffic, or exposure.",
        "PMF, expectation, variance, assumptions, derivations, approximations, and interview traps.",
        "Research engineer angle: modeling assumptions, likelihoods, gradients, simulations, rare-event behavior, and evaluation under distribution shift.",
    ])

    add_heading(doc, "2. Quick Selection Map", 1)
    add_table(
        doc,
        ["Business question", "Distribution", "Random variable", "Key parameter"],
        [
            ["Did a user click?", "Bernoulli", "X in {0, 1}", "p = click probability"],
            ["How many users clicked out of n?", "Binomial", "X = count of successes", "n, p"],
            ["How many attempts until first approval?", "Geometric", "X = trial number of first success", "p"],
            ["How many fraud alerts per hour?", "Poisson", "X = event count in interval", "lambda"],
        ],
        [2500, 1900, 3000, 1960],
    )

    add_heading(doc, "3. PMF: Probability Mass Function", 1)
    add_body(doc, "For discrete random variables, the PMF gives the probability that the variable takes an exact value. It must satisfy two conditions:")
    add_formula(doc, "1) P(X = x) >= 0 for every possible x")
    add_formula(doc, "2) Sum over all possible x of P(X = x) = 1")
    add_body(doc, "Intuition: the PMF is a probability budget. Every possible outcome gets some probability, and the total budget must be exactly 1.")

    add_heading(doc, "4. Bernoulli Distribution", 1)
    add_body(doc, "Use Bernoulli when there is exactly one binary event: success/failure, click/no click, fraud/not fraud, default/no default, churn/not churn.")
    add_formula(doc, "X ~ Bernoulli(p), where X = 1 with probability p and X = 0 with probability 1 - p")
    add_formula(doc, "P(X = x) = p^x (1 - p)^(1 - x), for x in {0, 1}")
    add_formula(doc, "E[X] = p")
    add_formula(doc, "Var(X) = p(1 - p)")
    add_heading(doc, "Why E[X] = p", 2)
    add_formula(doc, "E[X] = 0 * P(X = 0) + 1 * P(X = 1) = 0 * (1 - p) + 1 * p = p")
    add_heading(doc, "Why Var(X) = p(1 - p)", 2)
    add_formula(doc, "Var(X) = E[X^2] - (E[X])^2")
    add_formula(doc, "Since X is 0 or 1, X^2 = X, so E[X^2] = E[X] = p")
    add_formula(doc, "Var(X) = p - p^2 = p(1 - p)")
    add_callout(doc, "Interview trap:", "A Bernoulli variable is not only a coin toss. Any single binary label can be Bernoulli if the probability p is meaningful.")

    add_heading(doc, "5. Binomial Distribution", 1)
    add_body(doc, "Use Binomial when you count successes in a fixed number n of independent trials, each with the same probability p.")
    add_formula(doc, "X ~ Binomial(n, p)")
    add_formula(doc, "P(X = k) = C(n, k) p^k (1 - p)^(n - k), for k = 0, 1, ..., n")
    add_formula(doc, "E[X] = np")
    add_formula(doc, "Var(X) = np(1 - p)")
    add_heading(doc, "Why the PMF has C(n, k)", 2)
    add_body(doc, "One exact sequence with k successes and n-k failures has probability p^k(1-p)^(n-k). But the successes can appear in many positions. The number of ways to choose the k success positions is C(n, k).")
    add_formula(doc, "C(n, k) = n! / (k! (n-k)!)")
    add_heading(doc, "Why E[X] = np", 2)
    add_body(doc, "Represent the count as a sum of Bernoulli indicators.")
    add_formula(doc, "X = I1 + I2 + ... + In, where Ii ~ Bernoulli(p)")
    add_formula(doc, "E[X] = E[I1] + ... + E[In] = p + ... + p = np")
    add_heading(doc, "Why Var(X) = np(1-p)", 2)
    add_formula(doc, "If trials are independent: Var(X) = Var(I1 + ... + In) = Sum Var(Ii) = n p(1-p)")
    add_callout(doc, "Business interpretation:", "In A/B testing, if 10,000 users see a variant and each user converts with probability p, the number of conversions is Binomial under the fixed-n independent-user assumption.")

    add_heading(doc, "6. Geometric Distribution", 1)
    add_body(doc, "Use Geometric when you want the waiting time until the first success. Example: number of calls until first sale, number of transactions until first fraud, or number of attempts until a model finds a valid candidate.")
    add_formula(doc, "X ~ Geometric(p), where X = trial number of first success")
    add_formula(doc, "P(X = k) = (1 - p)^(k - 1) p, for k = 1, 2, 3, ...")
    add_formula(doc, "E[X] = 1 / p")
    add_formula(doc, "Var(X) = (1 - p) / p^2")
    add_heading(doc, "Why the PMF looks like this", 2)
    add_body(doc, "To get the first success on trial k, the first k-1 trials must fail and the kth trial must succeed.")
    add_formula(doc, "P(first success at k) = P(fail)^(k-1) P(success) = (1-p)^(k-1)p")
    add_heading(doc, "Memoryless property", 2)
    add_formula(doc, "P(X > s + t | X > s) = P(X > t)")
    add_body(doc, "Intuition: if every trial has the same success probability and trials are independent, previous failures do not make the next trial more likely to succeed. This is a strong assumption and often false in real systems with learning, fatigue, or adaptive targeting.")

    add_heading(doc, "7. Poisson Distribution", 1)
    add_body(doc, "Use Poisson for event counts in an interval when events happen independently, the average rate is stable, and simultaneous events are unlikely. It is common for rare events such as fraud alerts, server failures, claim arrivals, or customer complaints.")
    add_formula(doc, "X ~ Poisson(lambda)")
    add_formula(doc, "P(X = k) = e^(-lambda) lambda^k / k!, for k = 0, 1, 2, ...")
    add_formula(doc, "E[X] = lambda")
    add_formula(doc, "Var(X) = lambda")
    add_heading(doc, "How Poisson emerges from Binomial", 2)
    add_body(doc, "If there are many small opportunities for an event and each event is rare, Binomial(n, p) approaches Poisson(lambda) when n is large, p is small, and lambda = np.")
    add_formula(doc, "Binomial(n, p) with n -> infinity, p -> 0, np = lambda => Poisson(lambda)")
    add_heading(doc, "Research engineer warning: overdispersion", 2)
    add_body(doc, "Poisson assumes mean equals variance. Real data often has variance larger than the mean because of clustering, seasonality, user heterogeneity, or hidden state. If Var(X) >> E[X], consider Negative Binomial, zero-inflated models, or hierarchical modeling.")

    add_heading(doc, "8. Assumptions Checklist", 1)
    add_table(
        doc,
        ["Distribution", "Assumptions", "Common failure mode"],
        [
            ["Bernoulli", "One binary event; p is meaningful.", "Outcome is not truly binary or p changes by segment."],
            ["Binomial", "Fixed n, independent trials, same p.", "Users influence each other; p varies across cohorts."],
            ["Geometric", "Independent repeated trials, same p, first success matters.", "System learns from failures; attempts are not identical."],
            ["Poisson", "Events independent, rate stable, rare in tiny intervals.", "Seasonality, bursts, fraud rings, outages, or correlated arrivals."],
        ],
        [1800, 4200, 3360],
    )

    add_heading(doc, "9. Likelihood Thinking for Research Engineers", 1)
    add_body(doc, "Research engineers often need to connect probability distributions to model training. A likelihood asks: given parameter values, how likely is the observed data?")
    add_heading(doc, "Bernoulli likelihood", 2)
    add_formula(doc, "For observations x1, ..., xn in {0,1}: L(p) = Product p^xi (1-p)^(1-xi)")
    add_formula(doc, "Log-likelihood: l(p) = Sum [xi log p + (1-xi) log(1-p)]")
    add_body(doc, "This is the mathematical base of binary cross-entropy used in logistic regression and neural network binary classification.")
    add_heading(doc, "Binomial likelihood", 2)
    add_formula(doc, "If k successes out of n: L(p) proportional to p^k (1-p)^(n-k)")
    add_formula(doc, "MLE: p_hat = k / n")
    add_heading(doc, "Poisson likelihood", 2)
    add_formula(doc, "For counts x1, ..., xn: l(lambda) = Sum[-lambda + xi log(lambda) - log(xi!)]")
    add_formula(doc, "MLE: lambda_hat = sample mean")
    add_body(doc, "This is useful for count models, traffic modeling, demand forecasting, incident monitoring, and anomaly detection baselines.")

    add_heading(doc, "10. How BFSI Gets Help From These Distributions", 1)
    add_body(doc, "BFSI teams use discrete distributions because many financial decisions are about binary outcomes, event counts, and waiting time. The math helps convert uncertain customer or transaction behavior into measurable risk, monitoring rules, and business action.")
    add_table(
        doc,
        ["BFSI problem", "Distribution help", "Decision it supports"],
        [
            ["Loan default risk", "A borrower defaulting or not defaulting can be modeled as a Bernoulli outcome. A loan portfolio's number of defaults can start as Binomial when borrowers are treated as a homogeneous segment.", "Approve, reject, price risk, ask for collateral, or set credit limits."],
            ["Fraud detection", "Each transaction can be fraud/not fraud as Bernoulli. Fraud alerts per hour/day can be modeled with Poisson, but bursts may reveal fraud rings or attacks.", "Trigger alerts, tune thresholds, allocate fraud analysts, and detect abnormal spikes."],
            ["Insurance claims", "Claim/no-claim can be Bernoulli. Number of claims per policy group or time period can be Poisson-like if events are independent and rate is stable.", "Price premiums, reserve capital, detect unusual claim clusters, and plan claim operations."],
            ["Collections and recovery", "Successful recovery call/no success is Bernoulli. Number of attempts until first successful payment can use Geometric thinking.", "Choose contact strategy, estimate effort, prioritize accounts, and manage collection cost."],
            ["Operational risk", "Failed transaction counts, outage incidents, ATM failures, and complaint counts can be modeled as count processes.", "Set control limits, investigate process failures, and define service-level risk indicators."],
            ["Credit-card campaigns", "Response/no response is Bernoulli. Total responses among targeted customers can be Binomial.", "Estimate campaign ROI, run A/B tests, and choose customer segments for offers."],
        ],
        [2200, 3700, 3460],
    )
    add_heading(doc, "BFSI intuition", 2)
    add_bullets(doc, [
        "Bernoulli helps answer: What is the probability this customer, transaction, or policy becomes risky?",
        "Binomial helps answer: Out of this portfolio or campaign population, how many risky or successful cases should we expect?",
        "Geometric helps answer: How long or how many attempts until the first important event occurs?",
        "Poisson helps answer: Is the number of events in this time window normal, or is it unusually high?",
    ])
    add_heading(doc, "Where professionals must be careful", 2)
    add_bullets(doc, [
        "Customer risks are rarely identical. Segment by income, credit score, geography, tenure, product type, and behavior before assuming one common p.",
        "Defaults and fraud events are often correlated during macro stress, cyberattacks, or fraud-ring activity. Independence can fail badly.",
        "Poisson models can underestimate risk when event counts are bursty. Check whether variance is much larger than the mean.",
        "A model probability is not automatically a business decision. BFSI decisions need expected loss, regulatory constraints, fairness, explainability, and cost of false positives/false negatives.",
    ])

    add_heading(doc, "11. MAANG and BFSI Examples", 1)
    add_bullets(doc, [
        "Ads ranking: click/no-click is Bernoulli; total clicks among impressions is Binomial; rare ad policy violations can be Poisson-like.",
        "Search quality: success can mean user reformulates no query after result click; counts over sessions can use Binomial or Poisson models.",
        "Fraud monitoring: transaction flagged/not flagged is Bernoulli; number of fraud attempts per hour can be Poisson, but bursts may violate assumptions.",
        "Credit risk: default/no-default is Bernoulli at account level; portfolio defaults in a homogeneous segment may start as Binomial.",
        "Reliability engineering: number of incidents per week can be Poisson if rate is stable; outage clustering breaks the model.",
    ])

    add_heading(doc, "12. Interview Comparisons", 1)
    add_table(
        doc,
        ["Question", "Strong answer"],
        [
            ["Bernoulli vs Binomial", "Bernoulli is one binary trial. Binomial is the count of successes across n Bernoulli trials."],
            ["Binomial vs Poisson", "Binomial has fixed n and probability p. Poisson models counts over exposure with rate lambda, often for rare events."],
            ["Geometric vs Binomial", "Binomial asks how many successes in n trials. Geometric asks how many trials until the first success."],
            ["When not to use Poisson", "When events are clustered, variance is much larger than mean, or the rate changes across time/segments."],
        ],
        [2300, 7060],
    )

    add_heading(doc, "13. Practice Questions With Answers", 1)
    questions = [
        ("A landing page gets 1 conversion or no conversion from one visitor. Which distribution?", "Bernoulli, because one visitor gives one binary outcome."),
        ("Out of 1,000 visitors, how many convert if conversion probability is p?", "Binomial(1000, p), assuming fixed n, independent visitors, same p."),
        ("If default probability is 2 percent for each borrower in a homogeneous pool of 500, what is expected number of defaults?", "np = 500 * 0.02 = 10."),
        ("For Binomial(500, 0.02), what is variance?", "np(1-p) = 500 * 0.02 * 0.98 = 9.8."),
        ("Fraud alerts arrive at average 4 per hour. Probability of exactly 2 in an hour?", "Poisson: e^-4 * 4^2 / 2! = 8e^-4, about 0.1465."),
        ("If success probability per attempt is 0.2, expected attempts until first success?", "Geometric mean = 1/p = 5 attempts."),
        ("Why can Poisson be bad for fraud rings?", "Fraud rings create correlated bursts, violating independence and stable-rate assumptions."),
        ("What is the MLE for Bernoulli p after observing 70 successes in 200 trials?", "p_hat = 70/200 = 0.35."),
        ("A model predicts a probability for click/no-click. Which loss connects to Bernoulli likelihood?", "Binary cross-entropy / negative Bernoulli log-likelihood."),
        ("A count dataset has mean 3 and variance 20. What concern arises?", "Overdispersion; Poisson may understate uncertainty. Consider Negative Binomial or segmented/hierarchical models."),
    ]
    add_numbered(doc, [f"{q} Answer: {a}" for q, a in questions])

    add_heading(doc, "14. Research Engineer Extensions", 1)
    add_bullets(doc, [
        "Use simulations to validate intuition: draw Bernoulli trials, aggregate to Binomial, and compare empirical mean/variance with formulas.",
        "Check calibration: for Bernoulli predictions, among cases predicted near 0.2, roughly 20 percent should be positive.",
        "Use log-likelihood to compare model fit, but remember likelihood can improve while business utility worsens if decisions are poorly thresholded.",
        "Use residual checks for count models: if errors show bursts or time pattern, the iid Poisson assumption is probably broken.",
        "Connect distribution choice to data-generating process first, then choose model family. This is what separates research thinking from formula memorization.",
    ])

    add_heading(doc, "15. One-Page Revision Cheat Sheet", 1)
    add_table(
        doc,
        ["Distribution", "PMF", "Mean", "Variance", "Use when"],
        [
            ["Bernoulli(p)", "p^x(1-p)^(1-x)", "p", "p(1-p)", "One binary event"],
            ["Binomial(n,p)", "C(n,k)p^k(1-p)^(n-k)", "np", "np(1-p)", "Count successes in fixed n"],
            ["Geometric(p)", "(1-p)^(k-1)p", "1/p", "(1-p)/p^2", "Wait until first success"],
            ["Poisson(lambda)", "e^-lambda lambda^k/k!", "lambda", "lambda", "Count events per interval"],
        ],
        [1700, 2500, 1100, 1500, 2560],
    )

    add_heading(doc, "16. References and What to Read", 1)
    add_bullets(doc, [
        "BH: Introduction to Probability by Blitzstein and Hwang - distribution chapters, especially Bernoulli, Binomial, Geometric, and Poisson examples.",
        "PSDS: Practical Statistics for Data Scientists - probability distributions and business-facing statistical modeling intuition.",
        "ISLR: Use the Bernoulli likelihood connection when reading logistic regression and classification probability outputs.",
        "ESL: Advanced connection to likelihood, generalized linear models, and model diagnostics for non-Gaussian outcomes.",
    ])

    add_callout(
        doc,
        "Final memory hook:",
        "Bernoulli asks did it happen, Binomial asks how many happened out of fixed trials, Geometric asks how long until it happens, and Poisson asks how many happened in an interval."
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("Statistics Roadmap | Day 5 Discrete Distributions")
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = MUTED

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT.resolve())
