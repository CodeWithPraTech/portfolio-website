from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("Statistics_Day_5_Random_Variables_PMF_CDF_Expectation_Notes.docx")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 89, 89)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")


def configure_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    for style_name in ["List Bullet", "List Number"]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.add_run(text)
    return p


def add_formula(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(20, 38, 60)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    p = cell.paragraphs[0]
    run = p.add_run(title)
    run.bold = True
    run.font.color.rgb = DARK_BLUE
    p.add_run(f" {body}")
    doc.add_paragraph()


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, LIGHT_BLUE)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = DARK_BLUE
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    set_table_width(table, widths)
    doc.add_paragraph()


def build_doc():
    doc = Document()
    configure_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Statistics Day 5 Notes: Random Variables, PMF/CDF, Expectation")
    run.bold = True
    run.font.size = Pt(21)
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    sub = subtitle.add_run("Based on BH Ch. 3 | MAANG + BFSI + Research Engineer depth")
    sub.font.color.rgb = MUTED

    add_callout(
        doc,
        "Core idea:",
        "Random variables turn uncertain real-world outcomes into numbers. PMF/PDF/CDF describe their probability behavior. Expectation converts uncertainty into a long-run average used for risk, reward, and decision-making."
    )

    add_heading(doc, "1. Playlist Context", 1)
    add_table(
        doc,
        ["Field", "Day 5 value"],
        [
            ["Focus area", "Random Variables, PMF/CDF, Expectation"],
            ["Resource", "BH Ch. 3"],
            ["Important concepts", "Discrete and continuous random variables, PMF, PDF, CDF, expectation, linearity of expectation"],
            ["Practice", "Compute expected value for transaction loss, customer value, and product reward scenarios"],
        ],
        [2200, 7160],
    )

    add_heading(doc, "2. Random Variable", 1)
    add_body(doc, "A random variable is a function from outcomes to numbers. The uncertainty is in the outcome; the random variable gives a numerical measurement of that outcome.")
    add_formula(doc, "X: sample space -> real numbers")
    add_body(doc, "Example: toss a coin twice. The sample space is {HH, HT, TH, TT}. If X is the number of heads, then X(HH)=2, X(HT)=1, X(TH)=1, and X(TT)=0.")
    add_bullets(doc, [
        "X = 1 if a loan defaults, else 0.",
        "X = transaction loss amount.",
        "X = number of fraud alerts in one day.",
        "X = customer lifetime value.",
        "X = model loss on one prediction.",
    ])

    add_heading(doc, "3. Discrete vs Continuous Random Variables", 1)
    add_table(
        doc,
        ["Type", "Meaning", "Examples", "Main probability tool"],
        [
            ["Discrete", "Values are countable.", "0/1 fraud flag, number of clicks, number of defaults, number of complaints", "PMF and CDF"],
            ["Continuous", "Values can fall anywhere in an interval.", "transaction amount, stock return, time to default, claim amount", "PDF and CDF"],
        ],
        [1600, 2500, 3400, 1860],
    )

    add_heading(doc, "4. PMF: Probability Mass Function", 1)
    add_body(doc, "A PMF is used for a discrete random variable. It tells the probability of an exact value.")
    add_formula(doc, "p_X(x) = P(X = x)")
    add_formula(doc, "p_X(x) >= 0 for every x")
    add_formula(doc, "sum over all x of p_X(x) = 1")
    add_body(doc, "Example: X = number of heads in two fair coin tosses.")
    add_table(
        doc,
        ["x", "Outcomes", "P(X=x)"],
        [
            ["0", "TT", "1/4"],
            ["1", "HT, TH", "2/4"],
            ["2", "HH", "1/4"],
        ],
        [1200, 4200, 3960],
    )
    add_body(doc, "PMF answers exact questions such as: What is the probability of exactly 2 fraud alerts today?")

    add_heading(doc, "5. CDF: Cumulative Distribution Function", 1)
    add_body(doc, "A CDF gives the probability that the random variable is less than or equal to a threshold.")
    add_formula(doc, "F_X(x) = P(X <= x)")
    add_body(doc, "Using the same coin-toss example:")
    add_table(
        doc,
        ["x", "PMF P(X=x)", "CDF F(x)=P(X<=x)"],
        [
            ["0", "0.25", "0.25"],
            ["1", "0.50", "0.75"],
            ["2", "0.25", "1.00"],
        ],
        [1300, 3000, 5060],
    )
    add_bullets(doc, [
        "PMF asks: probability of exactly this value.",
        "CDF asks: probability up to this threshold.",
        "CDF is especially useful for risk limits, SLA limits, credit thresholds, and percentile analysis.",
    ])

    add_heading(doc, "6. PDF for Continuous Variables", 1)
    add_body(doc, "For continuous random variables, exact point probability is usually zero. A PDF gives density, not direct probability.")
    add_formula(doc, "P(X = x) = 0 for many continuous variables")
    add_formula(doc, "P(a <= X <= b) = area under the PDF from a to b")
    add_formula(doc, "P(a <= X <= b) = F(b) - F(a)")
    add_body(doc, "Example: for transaction amount X, asking P(X = INR 10,000 exactly) is usually not useful. A better question is P(9,000 <= X <= 11,000).")

    add_heading(doc, "7. Expectation", 1)
    add_body(doc, "Expectation is the long-run probability-weighted average of a random variable. It is written as E[X].")
    add_formula(doc, "Discrete: E[X] = sum over x of x * P(X = x)")
    add_formula(doc, "Continuous: E[X] = integral over x of x * f_X(x) dx")
    add_body(doc, "Expectation may not be a possible actual value. If expected purchases are 0.75, no individual customer makes 0.75 purchases. It means many similar customers average 0.75 purchases.")

    add_heading(doc, "8. Expectation Example: Purchases", 1)
    add_table(
        doc,
        ["Purchases x", "Probability", "x * P(X=x)"],
        [
            ["0", "0.50", "0.00"],
            ["1", "0.30", "0.30"],
            ["2", "0.15", "0.30"],
            ["3", "0.05", "0.15"],
            ["Total", "1.00", "0.75"],
        ],
        [1800, 2800, 4760],
    )
    add_formula(doc, "E[X] = 0(0.50) + 1(0.30) + 2(0.15) + 3(0.05) = 0.75")

    add_heading(doc, "9. Linearity of Expectation", 1)
    add_body(doc, "Linearity of expectation says the expectation of a sum is the sum of expectations.")
    add_formula(doc, "E[X + Y] = E[X] + E[Y]")
    add_formula(doc, "E[aX + b] = aE[X] + b")
    add_body(doc, "The powerful part: independence is not required for linearity of expectation. Independence matters for variance, but not for adding expectations.")
    add_body(doc, "If a bank has 3 loans with expected losses INR 1,000, INR 2,500, and INR 500, total expected loss is INR 4,000.")

    add_heading(doc, "10. Indicator Random Variables", 1)
    add_body(doc, "An indicator random variable is 1 if an event happens and 0 otherwise.")
    add_formula(doc, "I = 1 if event A happens, else 0")
    add_formula(doc, "E[I] = P(A)")
    add_body(doc, "This makes counting problems simple. If Xi = 1 when customer i defaults, then total defaults are X1 + X2 + ... + Xn. Expected total defaults are the sum of the default probabilities.")

    add_heading(doc, "11. BFSI Applications", 1)
    add_table(
        doc,
        ["BFSI use case", "Random variable", "Expectation/PMF/CDF use"],
        [
            ["Credit risk", "X = loss from a loan", "Expected loss supports approval, pricing, credit limit, and collateral decisions."],
            ["Fraud", "X = fraud loss if transaction is approved", "Expected fraud loss helps choose approve, OTP, review, or block."],
            ["Insurance", "X = claim amount", "Expectation supports premium pricing and reserve planning; CDF supports capital thresholds."],
            ["Collections", "X = recovery amount", "Expected recovery helps prioritize accounts and collection effort."],
            ["Operations", "X = failed transactions or downtime minutes", "PMF/CDF supports SLA planning, staffing, and control limits."],
        ],
        [2100, 2800, 4460],
    )

    add_heading(doc, "12. Research Engineer View", 1)
    add_body(doc, "In ML, model training is often about minimizing expected loss.")
    add_formula(doc, "Risk = E[L(Y, f(X))]")
    add_body(doc, "Because the real data-generating distribution is unknown, we approximate expected risk using training or validation data.")
    add_formula(doc, "Empirical risk = average observed loss on data")
    add_bullets(doc, [
        "Classification loss can be a random variable: 1 if prediction is wrong, 0 otherwise.",
        "Regression squared error is a random variable: (Y - prediction)^2.",
        "Recommendation reward is a random variable: click, purchase, retention, or revenue.",
        "Expected value is the bridge between probability theory and optimization objectives.",
    ])

    add_heading(doc, "13. Interview Traps", 1)
    add_bullets(doc, [
        "A random variable is not the outcome itself; it maps outcomes to numbers.",
        "PMF is for exact values of discrete random variables.",
        "CDF is cumulative: P(X <= x).",
        "PDF height is not probability; area under the PDF is probability.",
        "Expectation is a long-run average and may not be an actual possible value.",
        "Linearity of expectation does not require independence.",
    ])

    add_heading(doc, "14. Practice Questions With Answers", 1)
    add_numbered(doc, [
        "A customer makes 0, 1, 2, or 3 purchases with probabilities 0.4, 0.3, 0.2, 0.1. Expected purchases = 0(0.4)+1(0.3)+2(0.2)+3(0.1)=1.0.",
        "A loan has 3% default probability and INR 80,000 loss if default happens. Expected loss = 0.03 * 80,000 = INR 2,400.",
        "A transaction has 20% fraud probability and fraud loss INR 10,000. Manual review costs INR 200. Expected fraud loss = INR 2,000, so review is economically sensible.",
        "If X is number of failed payments, P(X <= 2) means probability that failed payments are at most 2.",
        "If X is transaction amount, P(X = 50,000) is usually not useful for a continuous variable. Use interval probability such as P(49,000 <= X <= 51,000).",
        "A portfolio has 500 loans, each with 4% default probability. Expected defaults = 500 * 0.04 = 20.",
        "E[X+Y] = E[X] + E[Y] helps portfolio risk because total expected loss is the sum of individual expected losses, even when losses are dependent.",
        "PMF example: P(number of fraud alerts today = 5). CDF example: P(total claim amount <= INR 100,000).",
    ])

    add_heading(doc, "15. Final Memory Hook", 1)
    add_table(
        doc,
        ["Concept", "Simple meaning"],
        [
            ["Random variable", "Turns uncertainty into a number."],
            ["PMF", "Probability of an exact discrete value."],
            ["CDF", "Probability up to a threshold."],
            ["PDF", "Density for continuous variables; area gives probability."],
            ["Expectation", "Long-run probability-weighted average."],
            ["Linearity", "Expected total equals sum of expected parts."],
        ],
        [2200, 7160],
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("Statistics Roadmap | Day 5 Random Variables")
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED

    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    build_doc()
