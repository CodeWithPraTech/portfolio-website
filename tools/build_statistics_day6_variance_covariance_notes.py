from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("Statistics_Day_6_Variance_Covariance_Correlation_Notes.docx")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 89, 89)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")


def configure_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    for style_name in ["List Bullet", "List Number"]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)


def add_body(doc, text):
    p = doc.add_paragraph()
    p.add_run(text)


def add_formula(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(20, 38, 60)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    p = cell.paragraphs[0]
    run = p.add_run(title)
    run.bold = True
    run.font.color.rgb = DARK_BLUE
    p.add_run(f" {body}")
    doc.add_paragraph()


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, LIGHT_BLUE)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = DARK_BLUE
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    set_table_width(table, widths)
    doc.add_paragraph()


def build_doc():
    doc = Document()
    configure_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Statistics Day 6 Notes: Variance, Covariance, Correlation")
    run.bold = True
    run.font.size = Pt(21)
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    sub = subtitle.add_run("Based on BH Ch. 4 + PSDS correlation sections | Research Engineer + Data Scientist depth")
    sub.font.color.rgb = MUTED

    add_callout(
        doc,
        "Core idea:",
        "Mean tells the center. Variance tells uncertainty around the center. Covariance tells how two variables move together. Correlation standardizes covariance so relationships can be compared across different units."
    )

    add_heading(doc, "1. Playlist Context", 1)
    add_table(
        doc,
        ["Field", "Day 6 value"],
        [
            ["Focus area", "Variance, Covariance, Correlation"],
            ["Resources", "BH Ch. 4; PSDS correlation sections"],
            ["Important concepts", "Variance, standard deviation, covariance, correlation, covariance matrix, correlation vs causation"],
            ["Practice", "Interpret correlation in credit score, income, default risk, product usage, model features, and portfolio risk examples"],
        ],
        [2200, 7160],
    )

    add_heading(doc, "2. Why Day 6 Matters", 1)
    add_body(doc, "Day 5 taught expectation, the average long-run value. Day 6 asks the next important question: how unstable is the value, and how do variables move together?")
    add_bullets(doc, [
        "For business: uncertainty changes decisions even when average outcome looks attractive.",
        "For BFSI: credit, market, fraud, and operational risk are about variability and co-movement.",
        "For data science: feature correlation affects interpretation, multicollinearity, model stability, and monitoring.",
        "For research engineering: variance appears in estimators, optimization noise, generalization, kernels, covariance matrices, and uncertainty modeling.",
    ])

    add_heading(doc, "3. Deviation From the Mean", 1)
    add_body(doc, "For a random variable X with mean mu = E[X], deviation means how far one outcome is from the mean.")
    add_formula(doc, "Deviation = X - mu")
    add_formula(doc, "E[X - mu] = E[X] - mu = 0")
    add_body(doc, "The average signed deviation is always zero because positive and negative deviations cancel. This is why variance squares deviations.")

    add_heading(doc, "4. Variance", 1)
    add_body(doc, "Variance measures the average squared distance from the mean. It is the most important mathematical measure of spread.")
    add_formula(doc, "Var(X) = E[(X - mu)^2]")
    add_formula(doc, "where mu = E[X]")
    add_formula(doc, "Equivalent shortcut: Var(X) = E[X^2] - (E[X])^2")
    add_body(doc, "The shortcut is often used in derivations, probability problems, and interview questions.")

    add_heading(doc, "4.1 Derivation of Var(X) = E[X^2] - E[X]^2", 2)
    add_formula(doc, "Var(X) = E[(X - mu)^2]")
    add_formula(doc, "= E[X^2 - 2muX + mu^2]")
    add_formula(doc, "= E[X^2] - 2muE[X] + mu^2")
    add_formula(doc, "= E[X^2] - 2mu^2 + mu^2")
    add_formula(doc, "= E[X^2] - mu^2")
    add_formula(doc, "Therefore Var(X) = E[X^2] - (E[X])^2")

    add_heading(doc, "4.2 Properties of Variance", 2)
    add_formula(doc, "Var(c) = 0")
    add_formula(doc, "Var(X + c) = Var(X)")
    add_formula(doc, "Var(aX) = a^2 Var(X)")
    add_formula(doc, "Var(aX + b) = a^2 Var(X)")
    add_formula(doc, "Var(X) >= 0")
    add_body(doc, "Adding a constant shifts all values but does not change spread. Multiplying by a scales spread, and because variance uses squared distance, the scaling becomes a squared.")

    add_heading(doc, "5. Standard Deviation", 1)
    add_body(doc, "Standard deviation is the square root of variance. It returns spread to the original unit of X.")
    add_formula(doc, "SD(X) = sqrt(Var(X))")
    add_table(
        doc,
        ["Metric", "Unit", "Use"],
        [
            ["Variance", "Squared unit", "Best for algebra, derivations, optimization, estimator analysis"],
            ["Standard deviation", "Original unit", "Best for interpretation and stakeholder communication"],
        ],
        [2200, 2200, 4960],
    )

    add_heading(doc, "6. Sample Variance vs Population Variance", 1)
    add_body(doc, "Population variance is the true variance of the data-generating process. Sample variance estimates it from observed data.")
    add_formula(doc, "Population variance: sigma^2 = (1/N) * sum from i=1 to N of (x_i - mu)^2")
    add_formula(doc, "Sample variance: s^2 = (1/(n-1)) * sum from i=1 to n of (x_i - x_bar)^2")
    add_body(doc, "The n-1 denominator is Bessel's correction. It compensates for the fact that x_bar is estimated from the same sample, which makes raw squared deviations too small on average.")

    add_heading(doc, "6.1 Why n - 1 Appears", 2)
    add_body(doc, "Once the sample mean is fixed, only n - 1 deviations are freely chosen because all deviations must sum to zero.")
    add_formula(doc, "sum from i=1 to n of (x_i - x_bar) = 0")
    add_body(doc, "So the sample loses one degree of freedom. Dividing by n - 1 gives an unbiased estimator of population variance under standard iid assumptions.")

    add_heading(doc, "7. Covariance", 1)
    add_body(doc, "Covariance measures whether two variables move together relative to their means.")
    add_formula(doc, "Cov(X, Y) = E[(X - mu_X)(Y - mu_Y)]")
    add_formula(doc, "Equivalent shortcut: Cov(X,Y) = E[XY] - E[X]E[Y]")
    add_table(
        doc,
        ["Covariance sign", "Meaning", "Example"],
        [
            ["Positive", "X above its mean tends to appear with Y above its mean.", "Higher income tends to appear with higher credit limit."],
            ["Negative", "X above its mean tends to appear with Y below its mean.", "Higher discount may appear with lower profit margin."],
            ["Near zero", "No strong linear co-movement.", "Feature and target may have no obvious linear relationship."],
        ],
        [1900, 3900, 3560],
    )

    add_heading(doc, "7.1 Covariance Derivation Shortcut", 2)
    add_formula(doc, "Cov(X,Y) = E[(X - mu_X)(Y - mu_Y)]")
    add_formula(doc, "= E[XY - Xmu_Y - Ymu_X + mu_Xmu_Y]")
    add_formula(doc, "= E[XY] - mu_YE[X] - mu_XE[Y] + mu_Xmu_Y")
    add_formula(doc, "= E[XY] - mu_Xmu_Y")
    add_formula(doc, "Therefore Cov(X,Y) = E[XY] - E[X]E[Y]")

    add_heading(doc, "7.2 Covariance Properties", 2)
    add_formula(doc, "Cov(X, X) = Var(X)")
    add_formula(doc, "Cov(aX + b, cY + d) = ac Cov(X, Y)")
    add_formula(doc, "Cov(X + Y, Z) = Cov(X, Z) + Cov(Y, Z)")
    add_formula(doc, "If X and Y are independent, Cov(X,Y) = 0")
    add_body(doc, "Important trap: zero covariance does not always mean independence. It only means no linear co-movement.")

    add_heading(doc, "8. Variance of a Sum", 1)
    add_body(doc, "Variance of a sum depends on individual variances and pairwise covariance.")
    add_formula(doc, "Var(X + Y) = Var(X) + Var(Y) + 2Cov(X,Y)")
    add_formula(doc, "Var(X - Y) = Var(X) + Var(Y) - 2Cov(X,Y)")
    add_body(doc, "If X and Y are independent, covariance is zero and the variance of the sum becomes the sum of variances.")
    add_formula(doc, "If independent: Var(X + Y) = Var(X) + Var(Y)")
    add_body(doc, "This is central in portfolio risk, ensembling, experimental measurement, and uncertainty propagation.")

    add_heading(doc, "9. Correlation", 1)
    add_body(doc, "Correlation standardizes covariance so it is unitless and bounded between -1 and 1.")
    add_formula(doc, "rho(X,Y) = Cov(X,Y) / (SD(X) * SD(Y))")
    add_formula(doc, "-1 <= rho <= 1")
    add_table(
        doc,
        ["Correlation", "Meaning", "Careful interpretation"],
        [
            ["+1", "Perfect positive linear relationship", "As X increases, Y increases exactly linearly."],
            ["0", "No linear relationship", "There may still be nonlinear dependence."],
            ["-1", "Perfect negative linear relationship", "As X increases, Y decreases exactly linearly."],
        ],
        [1700, 3300, 4360],
    )

    add_heading(doc, "9.1 Correlation Is Not Causation", 2)
    add_body(doc, "Correlation says two variables move together. It does not prove that one caused the other.")
    add_bullets(doc, [
        "Confounding: a third variable affects both variables.",
        "Reverse causality: Y may cause X instead of X causing Y.",
        "Selection bias: relationship appears only because of how data was collected.",
        "Time trend: two variables rise together over time without a direct causal link.",
        "Data leakage: feature appears correlated because it contains future information.",
    ])

    add_heading(doc, "10. Covariance Matrix", 1)
    add_body(doc, "For multiple variables, covariance is stored in a covariance matrix. Diagonal entries are variances; off-diagonal entries are covariances.")
    add_formula(doc, "Sigma_ij = Cov(X_i, X_j)")
    add_formula(doc, "Sigma = E[(X - mu)(X - mu)^T]")
    add_body(doc, "For two variables X1 and X2:")
    add_formula(doc, "Sigma = [[Var(X1), Cov(X1,X2)], [Cov(X2,X1), Var(X2)]]")
    add_bullets(doc, [
        "The covariance matrix is symmetric.",
        "It must be positive semi-definite.",
        "It appears in multivariate normal distributions, PCA, portfolio risk, Kalman filters, Gaussian processes, and uncertainty estimation.",
    ])

    add_heading(doc, "11. Research Engineer View", 1)
    add_body(doc, "Variance and covariance are not only descriptive statistics. They appear inside estimators, optimization, representation learning, kernels, and uncertainty modeling.")
    add_table(
        doc,
        ["Research topic", "Where Day 6 appears"],
        [
            ["Estimator variance", "A model or statistic is better when it has low variance without excessive bias."],
            ["Bias-variance tradeoff", "Expected test error decomposes into bias, variance, and irreducible noise."],
            ["Gradient noise", "Mini-batch gradients vary because each batch is a random sample."],
            ["PCA", "Principal components are eigenvectors of the covariance matrix."],
            ["Gaussian processes", "Covariance kernels define similarity and uncertainty between points."],
            ["Representation learning", "Correlation/covariance losses can decorrelate features or prevent collapse."],
        ],
        [2300, 7060],
    )

    add_heading(doc, "11.1 Bias-Variance Intuition", 2)
    add_body(doc, "A model's prediction changes if it is trained on a different sample. That instability is variance. A highly flexible model can have low training error but high variance.")
    add_formula(doc, "Expected test error = bias^2 + variance + irreducible noise")
    add_body(doc, "This is not the same variance as a single feature's variance, but it uses the same idea: instability around an expected value.")

    add_heading(doc, "12. Data Scientist View", 1)
    add_table(
        doc,
        ["Workflow stage", "How variance/covariance/correlation helps"],
        [
            ["EDA", "Detect spread, outliers, heavy tails, and unstable metrics."],
            ["Feature selection", "Find redundant correlated features and possible leakage."],
            ["Regression", "Multicollinearity inflates coefficient variance and makes explanations unstable."],
            ["Monitoring", "Feature drift can be seen through changing mean, variance, and correlation structure."],
            ["Experimentation", "Metric variance affects confidence intervals, power, and sample size."],
            ["Risk dashboards", "Volatility and correlation drive portfolio and market risk."],
        ],
        [2200, 7160],
    )

    add_heading(doc, "13. BFSI Applications", 1)
    add_table(
        doc,
        ["Use case", "Relevant concept", "Business decision"],
        [
            ["Credit risk", "Variance of default/loss", "Set credit limits, provisions, and capital buffers."],
            ["Fraud", "Correlation among merchant, amount, device, and location features", "Detect coordinated suspicious behavior and leakage risks."],
            ["Market risk", "Return variance and covariance", "Estimate volatility, diversify portfolio, compute risk limits."],
            ["Insurance", "Claim severity variance", "Set premiums, reserves, reinsurance, and stress buffers."],
            ["Collections", "Correlation between customer behavior and recovery", "Prioritize outreach and estimate expected recovery."],
        ],
        [1900, 2600, 4860],
    )

    add_heading(doc, "14. Worked Example: Variance", 1)
    add_body(doc, "Let X be the number of failed payments for a customer in a month.")
    add_table(
        doc,
        ["x", "P(X=x)", "x * P(X=x)", "x^2 * P(X=x)"],
        [
            ["0", "0.60", "0.00", "0.00"],
            ["1", "0.25", "0.25", "0.25"],
            ["2", "0.10", "0.20", "0.40"],
            ["3", "0.05", "0.15", "0.45"],
            ["Total", "1.00", "0.60", "1.10"],
        ],
        [1300, 2200, 2700, 3160],
    )
    add_formula(doc, "E[X] = 0.60")
    add_formula(doc, "E[X^2] = 1.10")
    add_formula(doc, "Var(X) = E[X^2] - E[X]^2 = 1.10 - 0.36 = 0.74")
    add_formula(doc, "SD(X) = sqrt(0.74) = 0.86 failed payments")
    add_body(doc, "Business interpretation: average failed payments are 0.60, but the standard deviation is 0.86, so customers vary meaningfully around the average.")

    add_heading(doc, "15. Worked Example: Covariance and Correlation", 1)
    add_body(doc, "Suppose X = standardized income score and Y = standardized credit limit score. If Cov(X,Y)=0.72, SD(X)=1.2, and SD(Y)=1.5:")
    add_formula(doc, "Corr(X,Y) = 0.72 / (1.2 * 1.5) = 0.72 / 1.8 = 0.40")
    add_body(doc, "Interpretation: income score and credit limit score have moderate positive linear association. This does not prove that income causes credit limit; policy rules, bank segment, geography, and credit history may confound the relationship.")

    add_heading(doc, "16. Common Interview Traps", 1)
    add_bullets(doc, [
        "High correlation does not imply causation.",
        "Zero correlation does not imply independence unless special assumptions hold, such as joint normality.",
        "Covariance is unit-dependent, so it is hard to compare across variable pairs.",
        "Correlation only measures linear association, not all types of dependence.",
        "Outliers can strongly distort covariance and Pearson correlation.",
        "Multicollinearity may not reduce predictive accuracy, but it can make coefficients unstable and hard to interpret.",
        "Variance of a sum includes covariance terms; ignoring correlation underestimates or overestimates risk.",
    ])

    add_heading(doc, "17. Pearson, Spearman, and Kendall", 1)
    add_table(
        doc,
        ["Measure", "Captures", "When to use"],
        [
            ["Pearson correlation", "Linear relationship", "Continuous variables with roughly linear relation and manageable outliers."],
            ["Spearman correlation", "Monotonic rank relationship", "Nonlinear monotonic patterns, ordinal variables, or outlier resistance."],
            ["Kendall tau", "Concordant vs discordant ranks", "Small samples, ordinal rankings, robust rank association."],
        ],
        [2200, 3300, 3860],
    )

    add_heading(doc, "18. Practice Questions With Answers", 1)
    add_numbered(doc, [
        "If Var(X)=25, SD(X)=5. Standard deviation is in the same unit as X.",
        "If Y=3X+10 and Var(X)=4, then Var(Y)=9*4=36. Adding 10 does not change variance.",
        "If Cov(X,Y)=0.30, SD(X)=2, SD(Y)=3, then Corr(X,Y)=0.30/(2*3)=0.05.",
        "If Corr(X,Y)=0, X and Y are not necessarily independent. They only have no linear correlation.",
        "If Var(X)=4, Var(Y)=9, Cov(X,Y)=2, then Var(X+Y)=4+9+2*2=17.",
        "If two portfolio assets have positive covariance, diversification benefit is weaker than if covariance is zero or negative.",
        "If credit score and income are highly correlated, both may be useful for prediction but coefficient interpretation in linear models can become unstable.",
        "If fraud score and manual review flag are strongly correlated, check for leakage because review may happen after suspicious behavior is already known.",
    ])

    add_heading(doc, "19. Interview-Ready Summary", 1)
    add_table(
        doc,
        ["Concept", "One-line answer"],
        [
            ["Variance", "Expected squared distance from the mean; measures spread."],
            ["Standard deviation", "Square root of variance; spread in original units."],
            ["Covariance", "Measures whether two variables move together relative to their means."],
            ["Correlation", "Standardized covariance between -1 and 1."],
            ["Covariance matrix", "Matrix containing all variances and covariances among features."],
            ["Correlation vs causation", "Correlation is association; causation needs identification, experiments, or causal assumptions."],
        ],
        [2300, 7060],
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("Statistics Roadmap | Day 6 Variance, Covariance, Correlation")
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED

    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    build_doc()
