from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "Statistics_Day_1_Notes.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(doc, text, level=1):
    paragraph = doc.add_heading(text, level=level)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return paragraph


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_callout(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_widths(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    p = cell.paragraphs[0]
    run = p.add_run(label)
    run.bold = True
    run.font.color.rgb = RGBColor(31, 77, 120)
    p.add_run(f" {text}")
    doc.add_paragraph()


def add_key_value_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Concept"
    table.rows[0].cells[1].text = "Interview-ready meaning"
    for cell in table.rows[0].cells:
        set_cell_shading(cell, "E8EEF5")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for concept, meaning in rows:
        cells = table.add_row().cells
        cells[0].text = concept
        cells[1].text = meaning
    set_table_widths(table, [1.8, 4.7])
    doc.add_paragraph()


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color in (
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(10 if name == "Heading 3" else 14)
        style.paragraph_format.space_after = Pt(6)


def build_doc():
    doc = Document()
    style_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Statistics Day 1 Notes")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Descriptive Statistics, EDA, Data Types, Outliers, and Interview Practice").italic = True

    add_callout(
        doc,
        "Day 1 goal:",
        "Learn how to inspect raw data, summarize it correctly, detect unusual values, and explain the business meaning in a MAANG-style interview.",
    )

    add_heading(doc, "1. Big Picture", 1)
    doc.add_paragraph(
        "Descriptive statistics and exploratory data analysis (EDA) help you understand what is happening in a dataset before you build a model. In simple terms, this is the first investigation of the data: what columns exist, what values are typical, how much values vary, whether anything is missing, and whether any unusual cases need attention."
    )
    add_callout(
        doc,
        "Interview line:",
        "Before modeling, I inspect data types, missing values, distributions, central tendency, spread, outliers, and group-level patterns so that the model is built on trustworthy assumptions.",
    )

    add_heading(doc, "2. Data Types and Data Type Ratio", 1)
    doc.add_paragraph(
        "A data type tells you what kind of information a column contains. The data type ratio tells you what percentage of columns belong to each type. This matters because it decides the right EDA, preprocessing, visualization, and modeling approach."
    )
    add_key_value_table(
        doc,
        [
            ("Numerical", "Numbers such as age, salary, transaction amount, credit score, or study hours. Focus on distribution, outliers, scaling, and correlation."),
            ("Categorical", "Labels such as city, product category, loan status, or user segment. Focus on frequency counts, rare categories, and encoding."),
            ("Ordinal", "Ordered categories such as Low, Medium, High or rating 1-5. Preserve the order while analyzing or encoding."),
            ("Boolean / Binary", "Two-value fields such as fraud/not fraud, clicked/not clicked, churned/not churned. Focus on class balance and conversion rates."),
            ("Datetime", "Dates or timestamps. Focus on trend, seasonality, recency, cohorts, and time leakage."),
        ],
    )
    doc.add_paragraph("Why the data type ratio is important:")
    add_key_value_table(
        doc,
        [
            ("Numerical-heavy dataset", "Expect more work on distributions, outliers, scaling, transformations, correlations, and regression-style thinking."),
            ("Categorical-heavy dataset", "Expect more work on frequency counts, rare categories, target leakage through IDs, one-hot/target encoding, and group comparisons."),
            ("Datetime-heavy dataset", "Expect more work on trend, seasonality, recency, cohort behavior, time windows, and avoiding future-data leakage."),
            ("Boolean-heavy dataset", "Expect more work on conversion rates, event flags, class balance, confusion matrix thinking, and rate comparisons."),
            ("Mixed dataset", "Use a combined strategy: numerical summaries, category frequency checks, time-based validation, and business segmentation."),
        ],
    )
    doc.add_paragraph("Questions to ask before calculating the ratio:")
    add_bullets(
        doc,
        [
            "How many columns are numerical, categorical, ordinal, boolean, and datetime?",
            "Which data type dominates the dataset, and what does that say about the business problem?",
            "Do ID-like categorical columns carry useful signal or create leakage risk?",
            "Are datetime columns true features, timestamps for ordering, or possible leakage sources?",
            "Will the dominant data type affect visualization, missing-value treatment, feature engineering, and model choice?",
        ],
    )
    doc.add_paragraph("Example data type ratio:")
    add_bullets(
        doc,
        [
            "Total columns = 20",
            "Numerical = 10 columns = 50%",
            "Categorical = 6 columns = 30%",
            "Datetime = 2 columns = 10%",
            "Boolean = 2 columns = 10%",
        ],
    )
    add_callout(
        doc,
        "Interview line:",
        "I check the data type ratio early because it tells me whether the main work will be numerical analysis, categorical encoding, time-based analysis, or class-balance checks.",
    )
    add_callout(
        doc,
        "MAANG-style relevance:",
        "Data type ratio shows that you understand the dataset before modeling. It helps you justify why you chose certain charts, preprocessing steps, validation strategy, and model families.",
    )

    add_heading(doc, "3. Central Tendency: Mean, Median, Mode", 1)
    add_key_value_table(
        doc,
        [
            ("Mean", "The average. Useful when data is fairly balanced, but sensitive to outliers."),
            ("Median", "The middle value after sorting. Better for skewed data such as salaries, house prices, and transaction amounts."),
            ("Mode", "The most frequent value. Useful for categorical values such as most common product category or city."),
        ],
    )
    doc.add_paragraph("Example: salaries = 4, 5, 6, 5, 100 lakhs. The mean becomes 24 lakhs, but the median is 5 lakhs. Median better represents the typical employee because one extreme salary pulls the mean upward.")
    add_callout(doc, "Interview line:", "If the data is skewed or has outliers, I prefer median over mean because median better represents the typical user or transaction.")

    add_heading(doc, "4. Spread: Variance, Standard Deviation, Range, IQR", 1)
    add_key_value_table(
        doc,
        [
            ("Variance", "Average squared distance from the mean. It measures spread, but the squared unit is less intuitive."),
            ("Standard deviation", "Typical distance from the mean. Low means consistent values; high means unstable or widely varying values."),
            ("Range", "Maximum minus minimum. Simple, but very sensitive to one extreme value."),
            ("IQR", "Q3 minus Q1. It captures the middle 50% of values and is robust to outliers."),
        ],
    )
    add_callout(doc, "Interview line:", "Mean tells me the center, but standard deviation and IQR tell me whether that center is reliable or whether the data varies a lot.")

    add_heading(doc, "5. Percentiles and Business Meaning", 1)
    doc.add_paragraph("A percentile tells you the value below which a percentage of observations fall. If the 90th percentile transaction amount is Rs. 2,000, then 90% of transactions are below Rs. 2,000.")
    add_bullets(
        doc,
        [
            "P50 is the median.",
            "P75 tells where the top 25% begins.",
            "P90/P95/P99 are useful for user spend, latency, fraud risk, and extreme behavior.",
        ],
    )
    add_callout(doc, "Business use:", "In product analytics and BFSI, percentiles are often more useful than averages because user behavior and money data are usually skewed.")

    add_heading(doc, "6. Outliers and Robust Statistics", 1)
    doc.add_paragraph("Outliers are unusual values. They may be errors, rare but valid behavior, fraud, VIP customers, system bugs, or important business events. Do not remove them blindly.")
    doc.add_paragraph("IQR outlier rule:")
    add_bullets(
        doc,
        [
            "Lower bound = Q1 - 1.5 x IQR",
            "Upper bound = Q3 + 1.5 x IQR",
            "Values outside this range are potential outliers.",
        ],
    )
    add_key_value_table(
        doc,
        [
            ("Less robust", "Mean, range, and variance change a lot when extreme values exist."),
            ("More robust", "Median, IQR, and percentiles stay more stable when extreme values exist."),
        ],
    )
    add_callout(doc, "Interview line:", "I investigate outliers before removing them because in fraud, credit risk, product analytics, and business data, outliers can be the signal.")

    add_heading(doc, "7. Practical EDA Workflow", 1)
    add_numbered(
        doc,
        [
            "Understand the business question and target metric.",
            "Inspect rows, columns, data types, and data type ratio.",
            "Check missing values and duplicate records.",
            "Summarize numerical columns using mean, median, standard deviation, IQR, and percentiles.",
            "Summarize categorical columns using frequency counts and rare-category checks.",
            "Visualize distributions and compare groups.",
            "Detect outliers and investigate their cause.",
            "Convert patterns into business hypotheses or modeling decisions.",
        ],
    )

    add_heading(doc, "8. MAANG-Style Interview Answer Template", 1)
    doc.add_paragraph(
        "First, I would inspect the data types, data type ratio, missing values, and duplicates. Then I would summarize numerical columns using mean, median, standard deviation, IQR, and percentiles. For categorical variables, I would check frequency counts and rare categories. I would visualize distributions to detect skewness and outliers. If outliers exist, I would investigate whether they are errors or meaningful business signals. Finally, I would compare metrics across important groups and convert findings into business hypotheses."
    )

    add_heading(doc, "9. Practice Questions With Answers", 1)
    qa = [
        ("What is the difference between mean and median?", "Mean is the average. Median is the middle value after sorting. Mean is affected by outliers; median is more stable."),
        ("When would median be better than mean?", "When data is skewed or has outliers, such as salaries, transaction amounts, house prices, or user spend."),
        ("Average salary is Rs. 30 LPA, but most employees earn Rs. 8 LPA. What could be happening?", "A few very high salaries are pulling the average upward. The distribution is likely right-skewed, so median and percentiles should be reported."),
        ("What does standard deviation tell you?", "It tells how much values typically vary around the mean. Low standard deviation means consistent values; high standard deviation means large variation."),
        ("Why is range not always a good measure of spread?", "Range uses only the minimum and maximum, so one extreme value can make the spread look misleadingly large."),
        ("What is IQR and why is it useful?", "IQR = Q3 - Q1. It measures the spread of the middle 50% and is useful because it is robust to outliers."),
        ("How do you detect outliers using the IQR method?", "Compute Q1, Q3, and IQR. Values below Q1 - 1.5 x IQR or above Q3 + 1.5 x IQR are potential outliers."),
        ("Should you always remove outliers?", "No. First check if they are errors, fraud, VIP customers, rare behavior, or system bugs. Then decide whether to keep, cap, remove, or analyze separately."),
        ("In a food delivery app, average order value increased suddenly. What EDA steps would you take?", "Compare mean vs median, inspect P90/P95/P99, check outliers, segment by city/category/campaign, and look for duplicates or data issues."),
        ("In fraud detection, why might outliers be important?", "Fraud often appears as unusual behavior: high amount, odd time, unusual location, or sudden spending spike. Outliers may be signal, not noise."),
        ("What summary statistics would you use for transaction amount data?", "Count, mean, median, standard deviation, min, max, P25, P50, P75, P90, P95, P99, and IQR."),
        ("If a dataset is highly skewed, which metrics would you report?", "Median, IQR, percentiles, and possibly a log-transformed distribution. Mean can be shown but with caution."),
        ("Difference between categorical and ordinal data?", "Categorical data has labels without order, such as city. Ordinal data has ordered labels, such as Low, Medium, High."),
        ("Give one example where mode is more useful than mean.", "Most common product category purchased. Mean does not make sense for category labels."),
        ("Explain EDA to a non-technical stakeholder.", "EDA is the first investigation of data to check what we have, whether it is clean, what patterns exist, and what unusual cases appear before making decisions or building models."),
    ]
    for question, answer in qa:
        p = doc.add_paragraph()
        q = p.add_run(f"Q: {question}")
        q.bold = True
        doc.add_paragraph(f"A: {answer}")

    add_heading(doc, "10. Mini Case", 1)
    doc.add_paragraph("Data: 2, 3, 4, 3, 5, 4, 3, 200")
    add_bullets(
        doc,
        [
            "Outlier: 200.",
            "Median is better than mean because 200 heavily increases the mean.",
            "Possible business reasons: user left app open, power user, bot, tracking bug, or special live event.",
            "Do not remove it immediately. Investigate first.",
            "Check user ID, timestamp, device, event logs, idle time, and whether similar long sessions exist.",
        ],
    )
    add_callout(doc, "Strong answer:", "I would not remove the 200-minute session blindly. I would first check if it is a tracking issue, idle session, bot behavior, or valid high-engagement user. Depending on business context, I may cap it, remove it, or analyze it separately.")

    add_heading(doc, "11. Day 1 Notes To Upload In Notion", 1)
    doc.add_paragraph(
        "Today I learned how to summarize data using data types, data type ratio, mean, median, mode, variance, standard deviation, percentiles, and IQR. I understood that mean and range are sensitive to outliers, while median, IQR, and percentiles are more robust. In business problems, outliers should be investigated before removal because they may represent fraud, VIP customers, bugs, or rare but important behavior. I also learned that EDA connects data quality and statistical summaries to business decisions before modeling."
    )

    doc.add_paragraph("References: Practical Statistics for Data Scientists, Chapter 1; OpenIntro Statistics review sections; personal MAANG/BFSI interview framing notes.")

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
